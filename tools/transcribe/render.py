"""Stage 4: render the canonical transcript into md / txt / docx / pdf.

A small `md_to_blocks` parser turns each page's Markdown into typed blocks
(heading / paragraph / list_item) that the docx and pdf writers consume. `md` is
written verbatim; `txt` strips Markdown syntax. The PDF is laid out with PyMuPDF
(no extra dependency) using the bundled DejaVuSans font (covers æ/ø/å).
"""
from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document

from core import ASSETS

FONT_FILE = ASSETS / "DejaVuSans.ttf"
FONT_NAME = "djv"

_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_ULIST = re.compile(r"^\s*[-*+]\s+(.*)$")
_OLIST = re.compile(r"^\s*(\d+)[.)]\s+(.*)$")


def md_to_blocks(md):
    """Parse a Markdown string into a flat list of typed blocks."""
    blocks = []
    para = []

    def flush():
        if para:
            blocks.append({"type": "paragraph", "text": " ".join(para).strip()})
            para.clear()

    for raw in (md or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            flush()
            continue
        m = _HEADING.match(line)
        if m:
            flush()
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            continue
        m = _OLIST.match(line)
        if m:
            flush()
            blocks.append({"type": "list_item", "ordered": True,
                           "marker": f"{m.group(1)}.", "text": m.group(2).strip()})
            continue
        m = _ULIST.match(line)
        if m:
            flush()
            blocks.append({"type": "list_item", "ordered": False,
                           "marker": "-", "text": m.group(1).strip()})
            continue
        para.append(line.strip())
    flush()
    return blocks


def transcript_to_markdown(transcript):
    """Join all page Markdown into one document string."""
    parts = [pg.get("markdown", "").strip() for pg in transcript["pages"]]
    return "\n\n".join(p for p in parts if p).strip() + "\n"


def _all_pages_blocks(transcript):
    return [md_to_blocks(pg.get("markdown", "")) for pg in transcript["pages"]]


# ---- writers -------------------------------------------------------------

def render_md(transcript, path):
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(transcript_to_markdown(transcript), encoding="utf-8")
    return str(out)


def render_txt(transcript, path):
    lines = []
    for blocks in _all_pages_blocks(transcript):
        for blk in blocks:
            if blk["type"] == "heading":
                lines.append(blk["text"])
                lines.append("")
            elif blk["type"] == "list_item":
                marker = blk["marker"] if blk["ordered"] else "-"
                lines.append(f"{marker} {blk['text']}")
            else:
                lines.append(blk["text"])
                lines.append("")
    text = "\n".join(lines).strip() + "\n"
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(text, encoding="utf-8")
    return str(out)


def render_docx(transcript, path):
    doc = Document()
    for i, blocks in enumerate(_all_pages_blocks(transcript)):
        if i > 0:
            doc.add_page_break()
        for blk in blocks:
            if blk["type"] == "heading":
                doc.add_heading(blk["text"], level=min(blk["level"], 9))
            elif blk["type"] == "list_item":
                style = "List Number" if blk["ordered"] else "List Bullet"
                doc.add_paragraph(blk["text"], style=style)
            else:
                doc.add_paragraph(blk["text"])
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


# PDF layout constants (A4, points)
PAGE_W, PAGE_H = 595.0, 842.0
MARGIN = 56.0
LINE_GAP = 1.35
_SIZES = {1: 18.0, 2: 15.0, 3: 13.0}
PARA_SIZE = 11.0
LIST_INDENT = 16.0


def _wrap(font, text, fs, max_w):
    words = text.split()
    if not words:
        return [""]
    lines, cur = [], words[0]
    for w in words[1:]:
        if font.text_length(cur + " " + w, fontsize=fs) <= max_w:
            cur += " " + w
        else:
            lines.append(cur)
            cur = w
    lines.append(cur)
    return lines


def render_pdf(transcript, path):
    font = fitz.Font(fontfile=str(FONT_FILE))
    doc = fitz.open()
    state = {"page": doc.new_page(width=PAGE_W, height=PAGE_H), "y": MARGIN}

    def draw_line(text, fs, x):
        lh = fs * LINE_GAP
        if state["y"] + lh > PAGE_H - MARGIN:
            state["page"] = doc.new_page(width=PAGE_W, height=PAGE_H)
            state["y"] = MARGIN
        state["page"].insert_text(
            (x, state["y"] + fs), text,
            fontname=FONT_NAME, fontfile=str(FONT_FILE), fontsize=fs,
        )
        state["y"] += lh

    for i, blocks in enumerate(_all_pages_blocks(transcript)):
        if i > 0:
            state["page"] = doc.new_page(width=PAGE_W, height=PAGE_H)
            state["y"] = MARGIN
        for blk in blocks:
            if blk["type"] == "heading":
                fs, x, prefix = _SIZES.get(blk["level"], PARA_SIZE), MARGIN, ""
            elif blk["type"] == "list_item":
                fs, x = PARA_SIZE, MARGIN + LIST_INDENT
                prefix = ("• " if not blk["ordered"] else f"{blk['marker']} ")
            else:
                fs, x, prefix = PARA_SIZE, MARGIN, ""
            text = prefix + blk["text"]
            for line in _wrap(font, text, fs, PAGE_W - MARGIN - x):
                draw_line(line, fs, x)
            state["y"] += fs * 0.45  # spacing between blocks

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out), garbage=4, deflate=True)
    doc.close()
    return str(out)


WRITERS = {"md": render_md, "txt": render_txt, "docx": render_docx, "pdf": render_pdf}


def render(transcript, out_dir, stem, formats):
    """Write `transcript` to each requested format. Returns {format: path}."""
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    written = {}
    for fmt in formats:
        if fmt not in WRITERS:
            raise ValueError(f"unknown format: {fmt!r} (choose from {sorted(WRITERS)})")
        written[fmt] = WRITERS[fmt](transcript, str(out / f"{stem}.{fmt}"))
    return written
